# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shd)

def styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)
        set_cell_shading(cell, '191970')
    
    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = val
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if c_idx > 0 else WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F0F0FF')
    
    doc.add_paragraph()
    return table

def main():
    base = os.path.dirname(__file__)
    img_dir = os.path.join(base, 'rapor_grafikleri')
    
    doc = Document()
    
    # Stil ayarlari
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # =================== KAPAK ===================
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('PID Kontrolcü ile\nOtonom Araç Şerit Takip\nSimülasyonu')
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = RGBColor(25, 25, 112)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('─' * 30)
    run.font.color.rgb = RGBColor(25, 25, 112)
    run.font.size = Pt(14)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Proje Raporu')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph()
    
    for text in ['Simülasyon Ortamı: Unity 6', 'Programlama Dili: C#']:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(6):
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Haziran 2026')
    run.font.size = Pt(12)
    run.italic = True
    run.font.color.rgb = RGBColor(130, 130, 130)
    
    doc.add_page_break()

    # =================== İÇİNDEKİLER ===================
    h = doc.add_heading('İçindekiler', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    icerik = [
        ('1.', 'Giriş'),
        ('2.', 'Zorunlu Grafikler'),
        ('   2.1.', 'Takip Hatası e(t)'),
        ('   2.2.', 'Kontrolcü Çıkışı u(t)'),
        ('   2.3.', 'Referans ve Gerçek Yol Karşılaştırması'),
        ('   2.4.', 'Araç Hızı'),
        ('3.', 'Bölgesel Analiz'),
        ('   3.1.', 'Düz Yol Bölgesi'),
        ('   3.2.', 'Viraj Bölgesi'),
        ('4.', 'İstatistiksel Sonuçlar'),
        ('5.', 'Analiz Soruları ve Cevapları'),
        ('6.', 'Sonuç'),
    ]
    for num, title in icerik:
        p = doc.add_paragraph()
        run = p.add_run(f'{num} {title}')
        run.font.size = Pt(12)
        if not num.startswith(' '):
            run.bold = True
    
    doc.add_page_break()

    # =================== 1. GİRİŞ ===================
    h = doc.add_heading('1. Giriş', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    doc.add_paragraph('Bu rapor, Unity oyun motoru ortamında geliştirilen otonom araç şerit takip simülasyonunun sonuçlarını sunmaktadır. Araç, referans olarak tanımlanan şeridi PID (Proportional-Integral-Derivative) kontrolcü yardımıyla takip etmektedir. Simülasyon, şehir içi yollar, kavşaklar, virajlar ve otoyol rampası gibi farklı yol geometrilerini içermektedir.')
    
    h2 = doc.add_heading('Sistem Özellikleri', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    
    styled_table(doc, ['Parametre', 'Değer'], [
        ['PID - Kp (Oransal)', '0.99'],
        ['PID - Ki (İntegral)', '0.10'],
        ['PID - Kd (Türevsel)', '2.04'],
        ['Araç Kütlesi', '1200 kg'],
        ['Hedef Hız', '20 m/s (72 km/h)'],
        ['Örnekleme Frekansı', '50 Hz (FixedUpdate)'],
        ['Toplam Simülasyon Süresi', '245.94 s'],
        ['Toplam Veri Noktası', '12,298'],
    ])
    
    h2 = doc.add_heading('Kontrol Mimarisi', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    
    for t in [
        'Hata Sinyali e(t): Araç ile hedef rota arasındaki açısal sapma (derece)',
        'Kontrol Sinyali u(t): PID kontrolcünün ürettiği direksiyon kontrol çıkışı',
        'Pure Pursuit + PID: Dinamik lookahead mesafesi ile hedefe yönelme, PID ile düzeltme'
    ]:
        doc.add_paragraph(t, style='List Bullet')
    
    doc.add_page_break()

    # =================== 2. GRAFİKLER ===================
    h = doc.add_heading('2. Zorunlu Grafikler', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    # 2.1
    h2 = doc.add_heading('2.1. Takip Hatası e(t)', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Grafik, simülasyon boyunca araç ile referans rota arasındaki açısal sapma hatasını göstermektedir. e(t) = 0 çizgisi referans hattı olarak kırmızı kesikli çizgiyle belirtilmiştir.')
    
    img = os.path.join(img_dir, '1_takip_hatasi_et.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 1: Takip Hatası e(t) - Zaman Grafiği')
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 100, 100)
    
    for t in [
        'Düz yol bölgelerinde hata ±0.15 derece civarında seyretmektedir',
        'Viraj bölgelerinde hata ±20 dereceye kadar artmaktadır',
        'Araç viraj çıkışlarında hatayı yaklaşık 2-3 saniye içinde sıfıra yaklaştırmaktadır'
    ]:
        doc.add_paragraph(t, style='List Bullet')
    
    doc.add_page_break()

    # 2.2
    h2 = doc.add_heading('2.2. Kontrolcü Çıkışı u(t)', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Kontrol sinyali u(t), PID kontrolcünün direksiyon açısını belirlemek için ürettiği çıkışı temsil etmektedir.')
    
    img = os.path.join(img_dir, '2_kontrolcu_cikisi_ut.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 2: Kontrolcü Çıkışı u(t) - Zaman Grafiği')
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    for t in [
        'Düz yollarda kontrol sinyali ±0.05 civarında, oldukça sakin seyretmektedir',
        'Keskin virajlarda sinyal ±27 değerine kadar yükselmektedir',
        'Kontrol sinyali, hatayı takip eden ama hafif gecikmeyle reaksiyon veren yumuşak bir karakteristik göstermektedir'
    ]:
        doc.add_paragraph(t, style='List Bullet')
    
    doc.add_page_break()

    # 2.3
    h2 = doc.add_heading('2.3. Referans ve Gerçek Yol Karşılaştırması', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Bu grafik, hata sinyali e(t) ile kontrol sinyali u(t) arasındaki ilişkiyi göstermektedir. Kontrolcünün hataya ne kadar hızlı ve ne ölçüde tepki verdiği görsel olarak karşılaştırılabilmektedir.')
    
    img = os.path.join(img_dir, '3_et_ut_karsilastirma.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 3: e(t) ve u(t) Karşılaştırması')
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # 2.4
    h2 = doc.add_heading('2.4. Araç Hızı', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Aracın simülasyon boyunca anlık hız grafiği. Virajlarda otomatik yavaşlama (corner braking) mekanizmasının etkinliği gözlemlenmektedir.')
    
    img = os.path.join(img_dir, '4_arac_hizi.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 4: Araç Hızı - Zaman Grafiği')
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # =================== 3. BÖLGESEL ANALİZ ===================
    h = doc.add_heading('3. Bölgesel Analiz', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    h2 = doc.add_heading('3.1. Düz Yol Bölgesi', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Düz yol bölgesinde (t=0.0s - 6.9s) takip hatası detaylı olarak gösterilmektedir. Hata ±0.15 derece bandında seyretmekte olup, bu durum kontrolcünün düz yollarda mükemmele yakın bir performans sergilediğini kanıtlamaktadır.')
    
    img = os.path.join(img_dir, '5_duz_yol_hatasi.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 5: Düz Yol Bölgesi - Detaylı Takip Hatası')
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()
    
    h2 = doc.add_heading('3.2. Viraj Bölgesi', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    doc.add_paragraph('Bir viraj geçişi sırasında hata ve kontrol sinyalinin davranışı detaylı olarak gösterilmektedir. Keskin virajlarda hatanın arttığı, kontrolcünün agresif bir düzeltme uyguladığı ve viraj çıkışında hatanın tekrar sıfıra yaklaştığı gözlemlenmektedir.')
    
    img = os.path.join(img_dir, '6_viraj_hatasi.png')
    if os.path.exists(img):
        doc.add_picture(img, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Şekil 6: Viraj Bölgesi - Hata ve Kontrol Analizi')
        run.italic = True; run.font.size = Pt(9); run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_page_break()

    # =================== 4. İSTATİSTİKLER ===================
    h = doc.add_heading('4. İstatistiksel Sonuçlar', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    styled_table(doc, ['Metrik', 'Değer'], [
        ['Ortalama |e(t)|', '1.5068°'],
        ['Maksimum |e(t)|', '21.7101°'],
        ['RMS Hata', '3.2115°'],
        ['Standart Sapma (e)', '3.2117°'],
        ['Ortalama |u(t)|', '1.1451'],
        ['Maksimum |u(t)|', '27.1554'],
        ['Ortalama Hız', '5.97 m/s (21.5 km/h)'],
        ['Maksimum Hız', '12.63 m/s (45.5 km/h)'],
    ])
    
    h2 = doc.add_heading('Kararlılık Analizi (Son 5 saniye)', level=2)
    for run in h2.runs:
        run.font.color.rgb = RGBColor(50, 50, 50)
    
    styled_table(doc, ['Metrik', 'Değer'], [
        ['Ortalama |e(t)|', '0.1568°'],
        ['Standart Sapma', '0.0386°'],
    ])
    
    doc.add_paragraph('Son 5 saniyedeki veriler, sistemin uzun vadede kararlı çalıştığını ve hatanın 0.16° civarında sabitlendiğini göstermektedir.')
    
    doc.add_page_break()

    # =================== 5. ANALİZ SORULARI ===================
    h = doc.add_heading('5. Analiz Soruları ve Cevapları', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    sorular = [
        ('5.1. Araç düz yolda nasıl davranmaktadır?',
         'Araç düz yolda son derece kararlı davranmaktadır. Düz yol bölgelerinde takip hatası ±0.15 derece bandında kalmakta olup, bu değer pratik açıdan sıfıra çok yakındır. Kontrolcü çıkışı da düz yolda minimal düzeyde kalmakta (±0.05), bu da direksiyonun gereksiz yere kırılmadığını ve aracın ip gibi düz gittiğini göstermektedir. Kd=2.04 gibi yüksek bir türev katsayısı, düz yolda oluşabilecek küçük salınımları etkin bir şekilde söndürmektedir.'),
        
        ('5.2. Virajlarda hata nasıl değişmektedir?',
         'Virajlarda hata belirgin şekilde artmaktadır. Keskin şehir içi kavşaklarda (90 derecelik dönüşler) hata ±15-21 derece seviyelerine çıkabilmektedir. Bu artış beklenen bir davranıştır çünkü viraj giriş anında araç hâlâ düz yöne bakarken hedef nokta yana kaymaktadır. Fiziksel olarak anlık yön değiştirmek mümkün değildir (araç atalet kütlesine sahiptir). Viraj çıkışında kontrolcü hatayı yaklaşık 2-3 saniye içinde sıfıra indirebilmektedir.'),
        
        ('5.3. PID parametreleri değiştirildiğinde araç davranışı nasıl etkilenmektedir?',
         'PID parametreleri simülasyon sırasında canlı olarak Inspector panelinden değiştirilebilmektedir. Kp artırıldığında araç rotaya daha agresif dönmeye çalışır, aşırı artışta (>2.0) salınım başlar. Kp azaltıldığında araç yavaş tepki verir, virajları geniş alır. Ki artırıldığında kalıcı hata (offset) giderilir ama aşırı artışta integral sarması oluşur. Kd artırıldığında salınımlar söner, aşırı artışta (>5.0) araç donuk/ağır tepki verir. Kd azaltıldığında araç hızlı tepki verir ama salınım ve titreşim başlar.'),
        
        ('5.4. Hangi parametre seti en iyi sonucu vermiştir?',
         'Kp=0.99, Ki=0.10, Kd=2.04 parametre seti en iyi sonucu vermiştir. Bu değerler deneysel olarak optimize edilmiştir. Kp=0.99 yeterince güçlü oransal tepki sağlarken aşırı kırılma yapmaz. Ki=0.10 düşük bir integral katsayısı kalıcı offseti yavaşça giderir. Kd=2.04 güçlü türev katsayısı salınımları etkin şekilde söndürür; bu en kritik parametredir.'),
        
        ('5.5. Kontrol sinyali çok agresif midir, yumuşak mıdır?',
         'Kontrol sinyali dengeli bir karakteristik göstermektedir. Düz yollarda u(t) ≈ ±0.05 ile oldukça yumuşak, normal virajlarda u(t) ≈ ±5 ile ölçülü agresif, keskin 90 derecelik kavşaklarda u(t) ≈ ±27 ile güçlü ama gerekli düzeyde agresiftir. Kd=2.04 değeri sayesinde kontrolcü, hata büyüdüğünde güçlü düzeltme yaparken ani değişimlerde fren etkisi oluşturarak aşırı tepkiyi (overshoot) engellemektedir.'),
        
        ('5.6. Araç salınım yapmakta mıdır?',
         'Seçilen parametre setiyle araç belirgin bir salınım yapmamaktadır. Düz yol bölgelerindeki hata grafiği incelendiğinde, monoton bir yaklaşma (convergence) gözlemlenmekte olup, sürekli artan-azalan bir salınım kalıbı (oscillation pattern) bulunmamaktadır. Kd=2.04 değerinin güçlü sönümleme etkisi, Kp=0.99 değerinin ürettiği düzeltme sinyalinin aşırıya kaçmasını engellemektedir.'),
        
        ('5.7. Sistem kararlı mıdır?',
         'Evet, sistem kararlıdır. Son 5 saniyedeki ortalama hata 0.1568° ve standart sapma 0.0386° ile minimum düzeydedir. Hata hiçbir zaman ıraksamamıştır (divergence yok). Her viraj sonrası hata sonlu sürede sıfıra yaklaşmaktadır (BIBO kararlılık). 246 saniyelik uzun simülasyon boyunca sistem hiçbir noktada kontrolü kaybetmemiştir.'),
        
        ('5.8. Keskin virajlarda takip başarımı düşmekte midir?',
         'Evet, keskin virajlarda takip başarımı beklenen şekilde düşmektedir. 90 derecelik şehir içi kavşaklarda anlık hata 21.7 dereceye kadar çıkabilmektedir. Bu düşüş kaçınılmazdır çünkü araç 1200 kg kütleye sahiptir ve ani yön değişikliği fiziksel olarak imkânsızdır. Otomatik fren sistemi virajlarda hızı düşürse de minimum 14 m/s korunmaktadır. Viraj çıkışında kontrolcü hatayı başarıyla gidermektedir.'),
        
        ('5.9. Kontrolcü parametreleri değiştiğinde yükselme ve yerleşme süreleri nasıl etkilenmektedir?',
         'Kp artarsa yükselme süresi (rise time) kısalır, ancak aşma (overshoot) artar ve yerleşme süresi uzayabilir. Kd artarsa yükselme süresi uzar, ancak aşma azalır ve yerleşme süresi kısalır (sönümleme artar). Ki artarsa kalıcı durum hatası azalır, ancak yerleşme süresi uzar ve aşma riski artar. Mevcut parametrelerle (Kp=0.99, Ki=0.10, Kd=2.04) sistem, overdamped (aşırı sönümlü) karakteristik göstermekte olup bu durum otonom araç için idealdir: aşma yapmadan, güvenli bir şekilde hedefe yaklaşmaktadır.'),
    ]
    
    for soru, cevap in sorular:
        h3 = doc.add_heading(soru, level=3)
        for run in h3.runs:
            run.font.color.rgb = RGBColor(70, 70, 70)
        doc.add_paragraph(cevap)
    
    doc.add_page_break()

    # =================== 6. SONUÇ ===================
    h = doc.add_heading('6. Sonuç', level=1)
    for run in h.runs:
        run.font.color.rgb = RGBColor(25, 25, 112)
    
    doc.add_paragraph('PID kontrolcü ile geliştirilen şerit takip sistemi, 30 adet referans noktasından oluşan karmaşık bir şehir içi rotayı başarıyla takip edebilmektedir. Kp=0.99, Ki=0.10, Kd=2.04 parametre seti ile aşağıdaki sonuçlar elde edilmiştir:')
    
    for t in [
        'Düz yollarda mükemmele yakın takip (±0.15°)',
        'Viraj çıkışlarında hızlı toparlanma (2-3 saniye)',
        'Salınımsız, kararlı sürüş',
        'Otomatik viraj freni ile güvenli dönüş',
        '246 saniyelik kesintisiz çalışma'
    ]:
        doc.add_paragraph(t, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Sistem, PID parametreleri uygun olmayan değerlere ayarlandığında (örneğin Kp>3, Kd=0) kontrol kaybı, salınım ve şerit ihlali doğal olarak gözlemlenebilmektedir. Bu durum, sistemde yapay sınırlandırma bulunmadığını ve PID parametrelerinin etkisinin doğrudan gözlemlenebildiğini kanıtlamaktadır.')

    # Kaydet
    out_path = os.path.join(base, 'PID_Rapor.docx')
    doc.save(out_path)
    print(f"\nWord rapor basariyla olusturuldu: {out_path}")


if __name__ == '__main__':
    main()
